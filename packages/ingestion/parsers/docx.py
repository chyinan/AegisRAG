from __future__ import annotations

from io import BytesIO
from zipfile import BadZipFile

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

from packages.ingestion.domain import (
    ParsedDocument,
    ParseRequest,
    Section,
    safe_title_from_filename,
)
from packages.ingestion.exceptions import EmptyDocumentContentError, GenericDocumentParseError

_MAX_HEADING_LEVEL = 9
_DOCX_PARSE_EXCEPTIONS = (
    BadZipFile,
    PackageNotFoundError,
    SyntaxError,
    KeyError,
    OSError,
    ValueError,
    TypeError,
    AttributeError,
)


class DocxParser:
    async def parse(self, request: ParseRequest) -> ParsedDocument:
        try:
            document = Document(BytesIO(request.content))
        except _DOCX_PARSE_EXCEPTIONS as exc:
            raise GenericDocumentParseError(details={"reason": "invalid_docx"}) from exc

        default_title = safe_title_from_filename(request.filename)
        title_root = default_title
        heading_stack: list[str | None] = [None] * _MAX_HEADING_LEVEL
        current_title_path = [default_title]
        current_lines: list[str] = []
        sections: list[Section] = []
        section_index = 1
        heading_count = 0

        def flush() -> None:
            nonlocal current_lines, section_index
            content = "\n".join(current_lines).strip()
            if not content:
                current_lines = []
                return
            title_path = list(current_title_path)
            sections.append(
                Section(
                    section_id=f"{request.version_id}:section-{section_index}",
                    tenant_id=request.tenant_id,
                    document_id=request.document_id,
                    version_id=request.version_id,
                    source_type="docx",
                    source_uri=request.source_uri,
                    title=title_path[-1],
                    title_path=title_path,
                    content=content,
                    page_start=None,
                    page_end=None,
                    acl=request.acl,
                    metadata={
                        "source_uri": request.source_uri,
                        "filename": request.filename,
                        "checksum": request.checksum,
                        "title_path": title_path,
                        "page_metadata": "unavailable",
                        "content_char_count": len(content),
                        "metadata": dict(request.metadata),
                    },
                )
            )
            section_index += 1
            current_lines = []

        try:
            paragraphs = list(document.paragraphs)
            for paragraph in paragraphs:
                text = paragraph.text.strip()
                style = paragraph.style
                style_name = (style.name if style is not None and style.name else "").strip()
                heading_level = _heading_level(style_name)

                if not text:
                    continue

                if _is_title_style(style_name):
                    flush()
                    title_root = text
                    heading_stack = [None] * _MAX_HEADING_LEVEL
                    current_title_path = [title_root]
                    heading_count += 1
                    continue

                if heading_level is not None:
                    flush()
                    heading_stack[heading_level - 1] = text
                    for index in range(heading_level, len(heading_stack)):
                        heading_stack[index] = None
                    current_title_path = [title_root] + [
                        item for item in heading_stack[:heading_level] if item
                    ]
                    heading_count += 1
                    continue

                current_lines.append(text)
        except _DOCX_PARSE_EXCEPTIONS as exc:
            raise GenericDocumentParseError(details={"reason": "docx_content_failed"}) from exc

        flush()

        if not sections:
            raise EmptyDocumentContentError()

        return ParsedDocument(
            tenant_id=request.tenant_id,
            document_id=request.document_id,
            version_id=request.version_id,
            source_type="docx",
            source_uri=request.source_uri,
            sections=sections,
            acl=request.acl,
            checksum=request.checksum,
            metadata={
                "filename": request.filename,
                "source_uri": request.source_uri,
                "section_count": len(sections),
                "heading_count": heading_count,
                "page_metadata": "unavailable",
            },
        )


def _is_title_style(style_name: str) -> bool:
    return style_name.casefold() == "title"


def _heading_level(style_name: str) -> int | None:
    normalized = style_name.casefold()
    if not normalized.startswith("heading "):
        return None
    level_text = normalized.removeprefix("heading ").strip()
    if not level_text.isdecimal():
        return None
    level = int(level_text)
    if 1 <= level <= _MAX_HEADING_LEVEL:
        return level
    return None
